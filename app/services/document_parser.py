"""文档解析服务 - 支持 PDF、DOCX、TXT、Markdown"""

import os
import chardet
from pathlib import Path


class DocumentParser:
    """解析多种文档格式，提取纯文本"""

    SUPPORTED_TYPES = {".pdf", ".docx", ".txt", ".md", ".markdown"}

    @staticmethod
    def parse(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return DocumentParser._parse_pdf(file_path)
        elif ext == ".docx":
            return DocumentParser._parse_docx(file_path)
        elif ext in (".txt", ".md", ".markdown"):
            return DocumentParser._parse_text(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        return "\n\n".join(pages)

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        from docx import Document
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n\n".join(paragraphs)

    @staticmethod
    def _parse_text(file_path: str) -> str:
        # 检测编码
        with open(file_path, "rb") as f:
            raw = f.read()
        detected = chardet.detect(raw)
        encoding = detected.get("encoding", "utf-8") or "utf-8"
        try:
            return raw.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            return raw.decode("utf-8", errors="replace")

    @staticmethod
    def is_supported(filename: str) -> bool:
        ext = Path(filename).suffix.lower()
        return ext in DocumentParser.SUPPORTED_TYPES
